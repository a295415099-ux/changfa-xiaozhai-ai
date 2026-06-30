from pathlib import Path
import json
import re
import zipfile
import xml.etree.ElementTree as ET

import pdfplumber
from docx import Document


SOURCE_DIR = Path("/Users/jayson/Downloads/竞品分析报告")
ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "tmp/competitor_reports"


def clean_text(text: str) -> str:
    text = re.sub(r"\r\n?", "\n", text or "")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def extract_docx(path: Path) -> str:
    doc = Document(path)
    parts = []
    for para in doc.paragraphs:
        value = para.text.strip()
        if value:
            parts.append(value)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " / ") for cell in row.cells]
            cells = [cell for cell in cells if cell]
            if cells:
                parts.append(" | ".join(cells))
    return clean_text("\n".join(parts))


def extract_pdf(path: Path) -> str:
    parts = []
    with pdfplumber.open(path) as pdf:
        for index, page in enumerate(pdf.pages, start=1):
            text = page.extract_text(x_tolerance=1, y_tolerance=3) or ""
            if text.strip():
                parts.append(f"--- Page {index} ---\n{text.strip()}")
    return clean_text("\n\n".join(parts))


def extract_pages(path: Path) -> tuple[str, list[str]]:
    found = []
    notes = []
    with zipfile.ZipFile(path) as zf:
        names = zf.namelist()
        notes.append("pages_archive_files=" + ",".join(names[:30]))
        for name in names:
            lower = name.lower()
            if lower.endswith((".txt", ".xml", ".html")):
                try:
                    data = zf.read(name)
                    text = data.decode("utf-8", errors="ignore")
                    text = re.sub(r"<[^>]+>", " ", text)
                    text = clean_text(text)
                    if len(text) > 80:
                        found.append(f"--- {name} ---\n{text}")
                except Exception as exc:
                    notes.append(f"failed_text:{name}:{exc}")
            elif lower.endswith(".iwa"):
                try:
                    data = zf.read(name)
                    text = data.decode("utf-8", errors="ignore")
                    # Pages stores most text in binary IWA protobuf chunks; this pulls readable Chinese/ASCII runs.
                    runs = re.findall(r"[\u4e00-\u9fffA-Za-z0-9，。、《》；：！？“”‘’（）()【】/\\-_%+,. ]{4,}", text)
                    runs = [clean_text(run) for run in runs if len(clean_text(run)) >= 4]
                    if runs:
                        found.append(f"--- {name} readable-runs ---\n" + "\n".join(runs))
                except Exception as exc:
                    notes.append(f"failed_iwa:{name}:{exc}")
    return clean_text("\n\n".join(found)), notes


def main() -> None:
    OUT.mkdir(parents=True, exist_ok=True)
    manifest = []
    for path in sorted(SOURCE_DIR.iterdir()):
        if path.name.startswith(".") or not path.is_file():
            continue
        suffix = path.suffix.lower()
        notes = []
        if suffix == ".docx":
            text = extract_docx(path)
            kind = "docx"
        elif suffix == ".pdf":
            text = extract_pdf(path)
            kind = "pdf"
        elif suffix == ".pages":
            text, notes = extract_pages(path)
            kind = "pages"
        else:
            continue
        out_path = OUT / f"{path.stem}.txt"
        out_path.write_text(text, encoding="utf-8")
        manifest.append({
            "name": path.name,
            "kind": kind,
            "source": str(path),
            "text_file": str(out_path),
            "chars": len(text),
            "notes": notes,
        })
    (OUT / "manifest.json").write_text(json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps(manifest, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
